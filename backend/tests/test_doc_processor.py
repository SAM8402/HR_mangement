"""Tests for doc_processor: text extraction, chunking, and parsing."""

from __future__ import annotations

import logging
from io import BytesIO
from unittest.mock import AsyncMock

import pytest
from docx import Document

from app.services.doc_processor import (
    _parse_text_locally,
    extract_text_from_docx,
    parse_docx,
    parse_role_text_with_llm,
)

logger = logging.getLogger(__name__)


# ── Fixtures ───────────────────────────────────────────────────────────────


@pytest.fixture
def docx_bytes():
    doc = Document()
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("Description")
    doc.add_paragraph("Build and maintain software products.")
    doc.add_paragraph("Responsibilities")
    doc.add_paragraph("Write clean code.")
    doc.add_paragraph("Review pull requests.")
    doc.add_paragraph("Skills")
    doc.add_paragraph("Python")
    doc.add_paragraph("Docker")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── _parse_text_locally ───────────────────────────────────────────────────


class TestParseTextLocally:
    def test_structured_text(self):
        lines = [
            "Backend Developer",
            "Description",
            "Builds and maintains server-side logic.",
            "Responsibilities",
            "Design APIs.",
            "Optimize database queries.",
            "Skills",
            "Python",
            "PostgreSQL",
        ]
        result = _parse_text_locally(lines)
        assert result["title"] == "Backend Developer"
        assert "server-side" in result["description"]
        assert "Design APIs" in result["responsibilities"]
        assert "Optimize database" in result["responsibilities"]
        assert result["skills"] == ["Python", "PostgreSQL"]

    def test_unstructured_text(self):
        lines = [
            "DevOps Engineer",
            "Manage cloud infrastructure and CI/CD pipelines.",
            "Deploy microservices.",
            "Monitor system health.",
        ]
        result = _parse_text_locally(lines)
        assert result["title"] == "DevOps Engineer"
        assert "cloud infrastructure" in result["description"]
        assert result["skills"] == []

    def test_empty_lines(self):
        result = _parse_text_locally([])
        assert result["title"] == "Untitled Role"
        assert result["description"] == ""
        assert result["skills"] == []

    def test_skills_keyword_variants(self):
        lines = [
            "Data Scientist",
            "Description",
            "Analyze data.",
            "Requirements",
            "Python",
            "SQL",
        ]
        result = _parse_text_locally(lines)
        assert result["skills"] == ["Python", "SQL"]

    def test_resp_keyword_variants(self):
        lines = [
            "Product Manager",
            "Description",
            "Own product roadmap.",
            "Key Responsibilities",
            "Define requirements.",
            "What you'll do",
            "Lead sprint planning.",
        ]
        result = _parse_text_locally(lines)
        assert "Define requirements" in result["responsibilities"]
        assert "Lead sprint planning" in result["responsibilities"]

    def test_colon_in_headers(self):
        lines = [
            "QA Engineer",
            "Description:",
            "Ensure product quality.",
            "Skills:",
            "Selenium",
            "Cypress",
        ]
        result = _parse_text_locally(lines)
        assert result["title"] == "QA Engineer"
        assert "product quality" in result["description"]
        assert result["skills"] == ["Selenium", "Cypress"]


# ── parse_docx ─────────────────────────────────────────────────────────────


class TestParseDocx:
    def test_parses_sections(self, docx_bytes):
        result = parse_docx(docx_bytes)
        assert result["title"] == "Software Engineer"
        assert "Build and maintain" in result["description"]
        assert "Write clean code" in result["responsibilities"]
        assert "Review pull requests" in result["responsibilities"]
        assert result["skills"] == ["Python", "Docker"]

    def test_empty_docx(self):
        doc = Document()
        buf = BytesIO()
        doc.save(buf)
        result = parse_docx(buf.getvalue())
        assert result["title"] == "Untitled Role"
        assert result["description"] == ""
        assert result["skills"] == []


# ── extract_text_from_docx ──────────────────────────────────────────────────


class TestExtractTextFromDocx:
    def test_extracts_raw_text(self, docx_bytes):
        text = extract_text_from_docx(docx_bytes)
        assert "Software Engineer" in text
        assert "Build and maintain" in text
        assert "Write clean code" in text
        assert "Python" in text

    def test_empty_docx(self):
        doc = Document()
        buf = BytesIO()
        doc.save(buf)
        text = extract_text_from_docx(buf.getvalue())
        assert text == ""



# ── parse_role_text_with_llm ────────────────────────────────────────────────


class TestParseRoleTextWithLlm:
    @pytest.mark.asyncio
    async def test_uses_llm_when_available(self, monkeypatch):
        from langchain_google_genai import ChatGoogleGenerativeAI
        from langchain_core.messages import AIMessage

        async def mock_ainvoke(self, input_data, *args, **kwargs):
            return AIMessage(
                content=(
                    '{"title": "DevOps Engineer", '
                    '"description": "Manage servers", '
                    '"responsibilities": "CI/CD", '
                    '"skills": ["Docker", "Linux"]}'
                )
            )

        monkeypatch.setattr(ChatGoogleGenerativeAI, "ainvoke", mock_ainvoke)

        result = await parse_role_text_with_llm("some job description text")
        assert result["title"] == "DevOps Engineer"
        assert result["description"] == "Manage servers"
        assert result["responsibilities"] == "CI/CD"
        assert result["skills"] == ["Docker", "Linux"]

    @pytest.mark.asyncio
    async def test_falls_back_to_local_parser_on_llm_failure(self, monkeypatch):
        from langchain_google_genai import ChatGoogleGenerativeAI

        async def mock_ainvoke_fail(self, input_data, *args, **kwargs):
            raise Exception("API quota exhausted")

        monkeypatch.setattr(ChatGoogleGenerativeAI, "ainvoke", mock_ainvoke_fail)

        text = (
            "Backend Developer\n"
            "Description\n"
            "Builds APIs.\n"
            "Responsibilities\n"
            "Design endpoints.\n"
            "Skills\n"
            "Python\n"
            "FastAPI"
        )
        result = await parse_role_text_with_llm(text)
        assert result["title"] == "Backend Developer"
        assert "Builds APIs" in result["description"]
        assert "Design endpoints" in result["responsibilities"]
        assert result["skills"] == ["Python", "FastAPI"]

    @pytest.mark.asyncio
    async def test_handles_unstructured_text_on_llm_failure(self, monkeypatch):
        from langchain_google_genai import ChatGoogleGenerativeAI

        async def mock_ainvoke_fail(self, input_data, *args, **kwargs):
            raise Exception("Rate limited")

        monkeypatch.setattr(ChatGoogleGenerativeAI, "ainvoke", mock_ainvoke_fail)

        text = "Just a simple note about a job."
        result = await parse_role_text_with_llm(text)
        assert result["title"] == "Just a simple note about a job."
        assert result["description"] == ""
        assert result["skills"] == []

    @pytest.mark.asyncio
    async def test_handles_malformed_llm_response(self, monkeypatch):
        from langchain_google_genai import ChatGoogleGenerativeAI
        from langchain_core.messages import AIMessage

        async def mock_ainvoke_bad(self, input_data, *args, **kwargs):
            return AIMessage(content="not valid json at all")

        monkeypatch.setattr(ChatGoogleGenerativeAI, "ainvoke", mock_ainvoke_bad)

        text = "Data Analyst\nDescription\nCrunch numbers.\nSkills\nExcel\nSQL"
        result = await parse_role_text_with_llm(text)
        # Should hit fallback because json.loads fails
        assert result["title"] == "Data Analyst"
        assert "Crunch numbers" in result["description"]
        assert result["skills"] == ["Excel", "SQL"]


# ── extract_text_from_pdf ──────────────────────────────────────────────────


class TestExtractTextFromPdf:
    def test_extracts_text(self):
        from io import BytesIO

        from pypdf import PdfWriter

        from app.services.doc_processor import extract_text_from_pdf

        writer = PdfWriter()
        writer.add_blank_page(612, 792)
        buf = BytesIO()
        writer.write(buf)

        text = extract_text_from_pdf(buf.getvalue())
        assert isinstance(text, str)
