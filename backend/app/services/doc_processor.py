"""Document text extraction and LLM-based parsing.

Provides utilities for extracting text from PDF and DOCX files,
chunking large texts, and using Gemini to extract structured role
information from raw job descriptions.
"""

from __future__ import annotations

import io
import json
import logging

from docx import Document
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_google_genai import ChatGoogleGenerativeAI
from pypdf import PdfReader

from app.core.config import settings

logger = logging.getLogger(__name__)


def extract_text_from_pdf(content: bytes) -> str:
    """Extract plain text from PDF bytes."""
    reader = PdfReader(io.BytesIO(content))
    logger.info("Parsing document (PDF, %d bytes)", len(content))
    text_parts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            text_parts.append(text)
    logger.info("Parsed %d pages from PDF document", len(reader.pages))
    return "\n\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    """Extract plain text from DOCX bytes."""
    doc = Document(io.BytesIO(content))
    logger.info("Parsing document (DOCX, %d bytes)", len(content))
    text_parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    logger.info("Parsed %d paragraphs from DOCX document", len(text_parts))
    return "\n".join(text_parts)


def _parse_text_locally(lines: list[str]) -> dict:
    """Extract structured role info from text lines using keyword detection."""
    title = lines[0] if lines else "Untitled Role"
    current_section = "description"
    section_buffers: dict[str, list[str]] = {
        "description": [],
        "responsibilities": [],
        "skills": [],
    }
    skills: list[str] = []

    skill_keywords = {"skills", "required skills", "requirements", "qualifications"}
    resp_keywords = {
        "responsibilities",
        "duties",
        "key responsibilities",
        "what you'll do",
    }

    for line in lines[1:]:
        lower = line.lower().rstrip(":")
        if lower in skill_keywords:
            current_section = "skills"
            continue
        if lower in resp_keywords:
            current_section = "responsibilities"
            continue

        if current_section == "skills":
            skill_text = line.strip()
            if skill_text:
                skills.append(skill_text)
        elif current_section == "responsibilities":
            section_buffers["responsibilities"].append(line)
        else:
            section_buffers["description"].append(line)

    description = " ".join(section_buffers["description"])
    responsibilities = " ".join(section_buffers["responsibilities"])

    if not description and not responsibilities:
        description = " ".join(lines[1:])

    return {
        "title": title or "Untitled Role",
        "description": description,
        "responsibilities": responsibilities,
        "skills": skills,
    }


def parse_docx(file_content: bytes) -> dict:
    """Parse a .docx file and extract structured role information locally."""
    doc = Document(io.BytesIO(file_content))
    logger.info("Parsing document with %d bytes", len(file_content))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    result = _parse_text_locally(paragraphs)
    logger.info(
        "Parsed %d sections from document %s", len(paragraphs), result["title"]
    )
    return result


async def parse_role_text_with_llm(text: str) -> dict:
    """Use Gemini to extract structured role info from raw text."""
    from app.ai.agent import get_llm

    llm = await get_llm(temperature=0.1)
    prompt = (
        "You are a helpful HR assistant. Parse the following job description text and extract:\n"
        "- title: The name of the role (e.g. Software Engineer)\n"
        "- description: A summary of the role\n"
        "- responsibilities: Key responsibilities\n"
        "- skills: A list of required skills (as a JSON array of strings)\n\n"
        "Return ONLY a valid JSON object with the keys 'title', 'description', 'responsibilities', 'skills'. "
        "Do not include markdown tags like ```json or similar in your response.\n\n"
        f"Text:\n{text}"
    )
    try:
        res = await llm.ainvoke(
            [
                SystemMessage(content="You parse text to JSON."),
                HumanMessage(content=prompt),
            ]
        )
        from app.ai.agent import clean_content

        cleaned = clean_content(res.content).strip()
        # Remove markdown code fence if present
        if cleaned.startswith("```"):
            cleaned = cleaned.split("\n", 1)[1]
        if cleaned.endswith("```"):
            cleaned = cleaned.rsplit("\n", 1)[0]
        if cleaned.startswith("json"):
            cleaned = cleaned.split("json", 1)[1].strip()

        data = json.loads(cleaned)
        return {
            "title": data.get("title", "Untitled Role"),
            "description": data.get("description", ""),
            "responsibilities": data.get("responsibilities", ""),
            "skills": data.get("skills", []),
        }
    except Exception as e:
        logger.error("LLM parsing failed, falling back to local parser: %s", e)
        return _parse_text_locally(text.split("\n"))
